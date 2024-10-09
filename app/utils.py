# app/utils.py
import asyncio
import io
import logging

import aiohttp
from models.condicionante import Condicionante
from models.licenca import Licenca
# from . import db
from bson.binary import Binary

import docx
import re
from datetime import datetime

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import pandas as pd

from io import BytesIO
import matplotlib.pyplot as plt

import matplotlib
from config import Config

# app/utils.py
# from app.db import db  # Importa a variável db do arquivo db.py
# DOCX_FIL

def carregar_dicionarios():
    Config.ARQUIVOS_NAO_ENTREGUES = {doc['_id']: doc for doc in Config.DB.nao_entregues.find()}
    Config.ARQUIVOS_ENTREGUES = {doc['_id']: doc for doc in Config.DB.entregues.find()}
    Config.ARQUIVOS_ENCERRADOS = {doc['_id']: doc for doc in Config.DB.encerrados.find()}

def salvar_dicionarios():

    # Salvando os arquivos não entregues
    if Config.ARQUIVOS_NAO_ENTREGUES:
        for k, v in Config.ARQUIVOS_NAO_ENTREGUES.items():
            # Verifica se o arquivo não está nas listas de entregues ou encerrados
            if k not in Config.ARQUIVOS_ENTREGUES and k not in Config.ARQUIVOS_ENCERRADOS:
                Config.DB.nao_entregues.update_one({'_id': k}, {'$set': v}, upsert=True)

    # Salvando os arquivos entregues e removendo da lista de "nao_entregues" se necessário
    if Config.ARQUIVOS_ENTREGUES:
        for k, v in Config.ARQUIVOS_ENTREGUES.items():
            # Verifica se o arquivo não está nas listas de não entregues ou encerrados
            if k not in Config.ARQUIVOS_NAO_ENTREGUES and k not in Config.ARQUIVOS_ENCERRADOS:
                Config.DB.entregues.update_one({'_id': k}, {'$set': v}, upsert=True)
            # Remove o arquivo da coleção "nao_entregues" se ele foi movido para "entregues"
            if k in Config.ARQUIVOS_NAO_ENTREGUES:
                Config.DB.nao_entregues.delete_one({'_id': k})
                # retirar também do Config.ARQUIVOS_NAO_ENTREGUES
                del Config.ARQUIVOS_NAO_ENTREGUES[k]

    # Salvando os arquivos encerrados e removendo das outras listas se necessário
    if Config.ARQUIVOS_ENCERRADOS:
        for k, v in Config.ARQUIVOS_ENCERRADOS.items():
            # Verifica se o arquivo não está nas listas de não entregues ou entregues
            if k not in Config.ARQUIVOS_NAO_ENTREGUES and k not in Config.ARQUIVOS_ENTREGUES:
                Config.DB.encerrados.update_one({'_id': k}, {'$set': v}, upsert=True)
            # Remove o arquivo da coleção "nao_entregues" ou "entregues" se ele foi movido para "encerrados"
            if k in Config.ARQUIVOS_NAO_ENTREGUES:
                Config.DB.nao_entregues.delete_one({'_id': k})
                # retirar também do Config.ARQUIVOS_NAO_ENTREGUES
                del Config.ARQUIVOS_NAO_ENTREGUES[k]
            if k in Config.ARQUIVOS_ENTREGUES:
                Config.DB.entregues.delete_one({'_id': k})
                # retirar também do Config.ARQUIVOS_ENTREGUES
                del Config.ARQUIVOS_ENTREGUES[k]


def listar_arquivos_pasta(drive_service, pasta_id):
    query = f"'{pasta_id}' in parents"
    results = drive_service.files().list(q=query, fields="files(id, name, mimeType)").execute()
    files = results.get('files', [])
    
    # Lista para armazenar todos os arquivos encontrados
    all_files = []
    
    # Verifica se há subpastas e busca recursivamente
    for file in files:
        if file['mimeType'] == 'application/vnd.google-apps.folder':
            sub_files = listar_arquivos_pasta(drive_service, file['id'])
            all_files.extend(sub_files)
        else:
            all_files.append(file)
    
    return all_files



def update_order(tabela, criterio):    
    if tabela == 'nao_entregues':
        Config.ORDEM_NAO_ENTREGUES = 'desc' if Config.ORDEM_NAO_ENTREGUES == 'asc' else 'asc'
        Config.CRITERIO_NAO_ENTREGUES = criterio
    elif tabela == 'entregues':
        Config.ORDEM_ENTREGUES = 'desc' if Config.ORDEM_ENTREGUES == 'asc' else 'asc'
        Config.CRITERIO_ENTREGUES = criterio
    elif tabela == 'encerradas':
        Config.ORDEM_ENCERRADOS = 'desc' if Config.ORDEM_ENCERRADOS == 'asc' else 'asc'
        Config.CRITERIO_ENCERRADOS = criterio

def filtrar_e_ordenar(arquivos, criterio, ordem, search_query, cor='todas'):
    # Filtra arquivos com base na consulta de busca
    arquivos_filtrados = {file_id: file_data for file_id, file_data in arquivos.items() if search_query.lower() in file_data['nome'].lower()}
    
    # Ordena arquivos filtrados
    criterios_data = ['data_carimbo', 'data_renovacao', 'data_vencimento', 'data_arquivamento']
    def get_key(item):
        value = item[1].get(criterio, '')
        if criterio in criterios_data and isinstance(value, str):
            try:
                return datetime.strptime(value, '%d/%m/%Y %H:%M:%S')
            except ValueError:
                try:
                    return datetime.strptime(value, '%d/%m/%Y')
                except ValueError:
                    return value
        return value
    if cor != 'todas':
        arquivos_filtrados = {file_id: file_data for file_id, file_data in arquivos_filtrados.items() if file_data['situacao_condicionantes'] == cor}

    return sorted(arquivos_filtrados.items(), key=get_key, reverse=(ordem == 'desc'))

def paginate(items, page, items_per_page):
    start = (page - 1) * items_per_page
    end = start + items_per_page
    return items[start:end]



def ordenar_arquivos(arquivos, criterio='numero', ordem='asc'):
    try:
        # Ordenar a lista de arquivos com base no critério e na ordem especificada
        arquivos_ordenados = sorted(arquivos, key=lambda x: x[1][criterio])
        
        # Inverter a lista se a ordem for descendente
        if ordem == 'desc':
            arquivos_ordenados = arquivos_ordenados[::-1]
        
        return arquivos_ordenados
    except Exception as e:
        # print(f"Erro ao ordenar arquivos: {e}")
        return []



# def atualizar_licenca_com_oficio(file_id, oficio):
#     # aqui vamos atualizar a licença, fazendo com que as condicionantes que tem no ofício sejam colocadas no lugar das que estão na licença e possuem o mesmo número

def encontrar_licenca_com_numero_de_processo(numero_processo):
    for file_id, file_data in Config.ARQUIVOS_ENTREGUES.items():
        if (file_data.get('numero_processo') == numero_processo):
            return file_id
    return None

def encontrar_licenca_com_numero_de_processo_encerradas(numero_processo):
    for file_id, file_data in Config.ARQUIVOS_ENCERRADOS.items():
        if (file_data.get('numero_processo') == numero_processo):
            return file_id
    return None

def encontrar_licenca_com_numero_de_processo_nao_entregues(numero_processo):
    for file_id, file_data in Config.ARQUIVOS_NAO_ENTREGUES.items():
        if (file_data.get('numero_processo') == numero_processo):
            return file_id
    return None

def verificar_se_oficio(file_id):
    # Verifica se o arquivo é um ofício
    if 'OF' in Config.ARQUIVOS_NAO_ENTREGUES[file_id]['tipo']:
        return True
    return False




def calcular_dias_prazo(prazo_str):
    # Verifica se o formato é MES/ANO
    match_mes_ano = re.match(r'(\d{2})/(\d{4})', prazo_str)
    if match_mes_ano:
        mes, ano = int(match_mes_ano.group(1)), int(match_mes_ano.group(2))
        prazo_data = datetime(ano, mes, 1)  # Primeiro dia do mês
        dias = (prazo_data - datetime.now()).days
        return dias

    # Verifica se o formato é N dias
    match_dias = re.match(r'(\d+)\s*\(\w+\)\s*dias', prazo_str)
    if match_dias:
        dias = int(match_dias.group(1))
        return dias

    return None

import docx
import re

def extrair_dados_oficio(docx_file):
    doc = docx.Document(docx_file)
    numero_processo_oficio = None
    condicionantes_oficio = []
    
    is_oficio = False
    current_condicionante = None
    current_numero = None
    current_prazos = []
    licenca_copia = None

    data = None

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # Verificar se o documento é um ofício
        if "OF/SEMMA" in text:
            is_oficio = True
            # print("OFICIO - Encontrado ofício")
        
        if is_oficio:
            # Extrair o número do processo
            if "Processo:" in text:
                numero_processo_oficio = text.split("Processo:")[-1].strip()
            
            # Verificar se estamos na seção de condicionantes
            match = re.match(r'Condicionante\s*n[°º]\s*(\d+)\s*-\s*(.*)', text, re.IGNORECASE)
            if match:
                current_numero = match.group(1)  # Captura o número da condicionante
                current_condicionante = match.group(2).strip()
                current_prazos = []  # Limpa os prazos anteriores para a nova condicionante

            # Verifica se a linha contém a palavra-chave "Prazo:" e extrai o(s) prazo(s)
            if "Prazo:" in text:
                prazos_texto = text.split("Prazo:")[-1].strip()
                prazos_lista = [p.strip().rstrip(";") for p in prazos_texto.split(",")]
                current_prazos.extend(prazos_lista)

            # Salvar a condicionante se houver
            if current_condicionante and current_numero:
                # print(f"OFICIO - Condicionante {current_numero}: {current_condicionante}")
                
                if not current_prazos:
                    current_prazos.append("Não Definido")
            
            # Condicionante atualizada
                # Busca a licença correspondente pelo número do processo
                for dicionario in [Config.ARQUIVOS_NAO_ENTREGUES, Config.ARQUIVOS_ENTREGUES, Config.ARQUIVOS_ENCERRADOS]:
                    for file_id, file_data in dicionario.items():
                        if file_data.get('numero_processo') == numero_processo_oficio:
                            # Carregar a licença
                            licenca = parse_licenca(file_data.get('data'))
                            if licenca and isinstance(licenca, Licenca):
                                # print("OFICIO - n da licenca original:", licenca.numero)
                                licenca_copia = licenca.copiar()  # Copia os atributos da licença original
                                
                                # Atualizar a condicionante
                                for cond in licenca_copia.condicionantes:
                                    if cond.numero.lstrip('0') == current_numero.lstrip('0'):
                                        cond.data_oficio = datetime.now().strftime('%d/%m/%Y')  # Atualiza a data do ofício
                                        cond.tem_oficio = True
                                        cond.prazo = current_prazos[0]  # Atualiza o prazo
                                        # print("OFICIO - Prazo da condicionante atualizado para:", cond.prazo)
                                        # print(cond.prazo)
                                        # condicionantes_oficio.append(cond)
                                        if (cond not in condicionantes_oficio):
                                            condicionantes_oficio.append(cond)
                                        
                                        # data = licenca_copia
                                        # print("Condicionante atualizada com sucesso")
                                        # salvar_dicionarios()
                                        break
                                licenca_copia.condicionantes = condicionantes_oficio
                            break

    if numero_processo_oficio:
        print(f'Número do Processo do Ofício: {numero_processo_oficio}')

    # Salvar os dicionários globalmente
    # salvar_dicionarios()
    
    if (licenca_copia is None):
        # print("extrair_dados_oficio: Licença não encontrada")
        return None
    licenca_copia.condicionantes = licenca_copia.condicionantes_unique()
    licenca_copia.tipo = "OF"

    return licenca_copia



def extrair_dados_licenca(docx_file):
    doc = docx.Document(docx_file)
    numero = None
    requerente = None
    cnpj = None
    numero_processo = None
    endereco = None
    atividade = None
    classe = None
    porte = None
    potencial_poluidor = None
    coordenadas = None
    validade = None
    condicionantes = []
    is_oficio = False
    # print("Extraindo dados da licença...")
    is_condicionante_section = False
    current_condicionante = None
    current_prazos = []
    current_numero = None  # Novo atributo para armazenar o número da condicionante

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue
        # Verificar se o documento é um ofício
        if "OF/SEMMA" in text:
            is_oficio = True
            # print("Ofício encontrado")
            data = extrair_dados_oficio(docx_file)
            break

        if 'REQUERENTE:' in text:
            requerente = text.split('REQUERENTE:')[-1].strip()
        if 'CPF/CNPJ:' in text:
            cnpj = text.split('CPF/CNPJ:')[-1].strip()
        if 'Processo nº' in text:
            numero_processo = text.split('Processo nº')[-1].split(',')[0].strip()
        if 'ENDEREÇO:' in text:
            endereco = text.split('ENDEREÇO:')[-1].strip()
        if 'ATIVIDADE:' in text:
            atividade = text.split('ATIVIDADE:')[-1].strip()
        if 'CLASSE:' in text:
            classe = text.split('CLASSE:')[-1].strip()[:3]
        if 'PORTE:' in text:
            porte = text.split('PORTE:')[-1].strip().split(' ')[0]
        if 'POTENCIAL POLUIDOR:' in text:
            potencial_poluidor = text.split('POTENCIAL POLUIDOR:')[-1].strip()
        if 'COORDENADAS GEOGRÁFICAS:' in text:
            coordenadas = text.split('COORDENADAS GEOGRÁFICAS:')[-1].strip()
        if 'VALIDADE:' in text:
            validade = text.split('VALIDADE:')[-1].strip()

        # Verificar se estamos na seção de condicionantes
        if "Condicionantes:" in text:
            is_condicionante_section = True
            continue

        if is_condicionante_section:
            # Identificar se o texto começa com um número seguido de um hífen, indicando uma nova condicionante
            match = re.match(r'(\d+)\s*-\s*(.*)', text)
            if match:
                # Se já temos uma condicionante em andamento, salvar a anterior antes de começar a nova
                if current_condicionante:
                    # Se não houver prazos, salvar como "Não Definido"
                    if not current_prazos:
                        current_prazos.append("Não Definido")

                    # Remover "Prazo:" da descrição, se presente
                    descricao_final = current_condicionante.split("Prazo:")[0].strip()

                    # Salva a condicionante atual, uma para cada prazo
                    for prazo in current_prazos:
                        # Verifica se a condicionante já não está na lista
                        condicionantes.append(Condicionante(descricao=descricao_final, prazo=prazo, numero=current_numero))

                # Iniciar uma nova condicionante
                current_numero = match.group(1)  # Captura o número da condicionante
                current_condicionante = match.group(2).strip()
                current_prazos = []  # Limpa os prazos anteriores para a nova

            # Verifica se a linha contém a palavra-chave "Prazo:" e extrai o(s) prazo(s)
            if "Prazo:" in text:
                prazos_texto = text.split("Prazo:")[-1].strip()
                prazos_lista = [p.strip().rstrip(";") for p in prazos_texto.split(",")]
                
                # Adiciona os prazos encontrados à lista
                current_prazos.extend(prazos_lista)

            # Se estamos dentro de uma condicionante e não encontramos uma nova, adicionamos mais texto à descrição
            elif current_condicionante:
                # Adiciona texto somente se a linha não indicar o início de uma nova condicionante
                if not re.match(r'^\d+\s*-\s*', text):
                    current_condicionante += ' ' + text.strip()

    if (is_oficio is True):
        # print("extrair_dados_licenca: Retornando ofício\n")
        # print(data)
        return data

    # Adicionar a última condicionante, se houver
    if current_condicionante:
        # Se não houver prazos, salva como "Não Definido"
        if not current_prazos:
            current_prazos.append("Não Definido")

        # Salvar a última condicionante
        for prazo in current_prazos:
            condicionantes.append(Condicionante(descricao=current_condicionante, prazo=prazo, numero=current_numero))

    # Verificar tabelas apenas para o número e tipo da licença
    padrao_licenca_tabela = re.compile(r'\b(\w+) N°\s*(\d{1,3}/\d{4})\b', re.IGNORECASE)
    tipo_licenca = None
    numero = None

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                
                if not text:
                    continue
                
                # Verificar se o texto corresponde ao padrão do tipo e número da licença
                match_licenca = padrao_licenca_tabela.search(text)
                if match_licenca:
                    tipo_licenca = match_licenca.group(1)  # Captura o tipo da licença (ex: "LO", "LP", etc.)
                    numero = match_licenca.group(2)  # Captura o número da licença (ex: "015/2022")

    # print(f'Número: {numero}')
    # print(f'Tipo da Licença: {tipo_licenca}')
    # print(f'Requerente: {requerente}')
    # print(f'CNPJ: {cnpj}')
    # print(f'Número do Processo: {numero_processo}')

    # Retorne a licença com o tipo também
    return Licenca(numero, tipo_licenca, requerente, cnpj, numero_processo, endereco, atividade, classe, porte, potencial_poluidor, coordenadas, validade, condicionantes)



# Prefixo para o Google Drive
async def processar_arquivo(session, file_id, file_data):
    # print(f'Processando arquivo {file_id}...')

    prefix = 'https://drive.google.com/uc?/export=download&id='
    if not file_data.get('data'):  # Verifica se os dados ainda não foram carregados
        try:
            # Verifica se o file_id já existe no banco de dados
            if Config.DOCX_FILES_COLLECTION.find_one({'file_id': file_id}):
                logging.info(f'Licença para {file_id} já está no banco de dados.')
                return

            url = f'{prefix}{file_id}'
            async with session.get(url) as resposta:
                resposta.raise_for_status()
                conteudo = await resposta.read()

                licenca = extrair_dados_licenca(io.BytesIO(conteudo))
                if licenca is not None and isinstance(licenca, Licenca):
                    # Salvar o arquivo .docx no MongoDB
                    Config.DOCX_FILES_COLLECTION.insert_one({
                        'file_id': file_id,
                        'data': Binary(conteudo)
                    })

                    file_data['data'] = repr(licenca)
                    file_data['numero_processo'] = licenca.numero_processo
                    file_data['numero'] = licenca.numero
                    file_data['tipo'] = licenca.tipo
                    file_data['nome'] = licenca.requerente
                    logging.info(f'Licença para {file_id} carregada e salva no banco de dados.')

                else:
                    logging.info(f'Erro ao carregar licença para {file_id}')

        except Exception as e:
            logging.error(f'Erro ao carregar licença para {file_id}: {e}')


async def carregar_licencas():
    async with aiohttp.ClientSession() as session:
        # print('Carregando licenças...')
        tasks = []
        for dicionario in [Config.ARQUIVOS_NAO_ENTREGUES, Config.ARQUIVOS_ENTREGUES, Config.ARQUIVOS_ENCERRADOS]:
            # print("aaa")
            for file_id, file_data in dicionario.items():
                tasks.append(processar_arquivo(session, file_id, file_data))
                licenca = parse_licenca(file_data.get('data'))
                if licenca and isinstance(licenca, Licenca):
                    file_data['numero_processo'] = licenca.numero_processo
                    file_data['numero'] = licenca.numero
                    file_data['tipo'] = licenca.tipo
                    file_data['nome'] = licenca.requerente
        
        # Execute todas as tarefas em paralelo
        await asyncio.gather(*tasks)
    
    salvar_dicionarios()

# Função para executar o loop assíncrono
def executar_carregar_licencas():
    asyncio.run(carregar_licencas())
    # print('Licenças carregadas com sucesso!')

import re

# Função que recebe uma string com os dados de uma licença e retorna um objeto Licenca
def parse_licenca(data):
    if data is None or not isinstance(data, str):
        return None

    numero = None
    tipo = None
    requerente = None
    cnpj = None
    numero_processo = None
    endereco = None
    atividade = None
    classe = None
    porte = None
    potencial_poluidor = None
    coordenadas = None
    validade = None
    condicionantes = []
    data_carimbo = None
    numero_ja_lido = False


    patterns = {
        'numero': re.compile(r"numero='([^']*)',"),
        'tipo': re.compile(r"tipo='([^']*)',"),
        'requerente': re.compile(r"requerente='([^']*)',"),
        'cnpj': re.compile(r"cnpj='([^']*)',"),
        'numero_processo': re.compile(r"numero_processo='([^']*)',"),
        'endereco': re.compile(r"endereco='([^']*)',"),
        'atividade': re.compile(r"atividade='([^']*)',"),
        'classe': re.compile(r"classe='([^']*)',"),
        'porte': re.compile(r"porte='([^']*)',"),
        'potencial_poluidor': re.compile(r"potencial_poluidor='([^']*)',"),
        'coordenadas': re.compile(r"coordenadas='([^']*)',"),
        'validade': re.compile(r"validade='([^']*)',"),
        'data_carimbo': re.compile(r"data_carimbo='([^']*)',")
    }

    for key, pattern in patterns.items():
        match = pattern.search(data)
        if match:
            if key == 'numero' and not numero_ja_lido:
                numero = match.group(1)
                numero_ja_lido = True
            elif key == 'tipo':
                tipo = match.group(1)
            elif key == 'requerente':
                requerente = match.group(1)
            elif key == 'cnpj':
                cnpj = match.group(1)
            elif key == 'numero_processo':
                numero_processo = match.group(1)
            elif key == 'endereco':
                endereco = match.group(1)
            elif key == 'atividade':
                atividade = match.group(1)
            elif key == 'classe':
                classe = match.group(1)
            elif key == 'porte':
                porte = match.group(1)
            elif key == 'potencial_poluidor':
                potencial_poluidor = match.group(1)
            elif key == 'coordenadas':
                coordenadas = match.group(1)
            elif key == 'validade':
                validade = match.group(1)
            elif key == 'data_carimbo':
                data_carimbo = match.group(1)

    condicionante_pattern = re.compile(
    #  agora vou pegar também o numero ao final
        r"Condicionante\(descricao='([^']*)', prazo='([^']*)', cumprida='([^']*)', tempo_restante='([^']*)', situacao='([^']*)', numero='([^']*)', tem_oficio='([^']*)', data_oficio='([^']*)', data_cumprimento='([^']*)'\)"
    )

    for match in condicionante_pattern.finditer(data):
        descricao = match.group(1)
        prazo = match.group(2)
        cumprida = match.group(3)
        tempo_restante = match.group(4)
        situacao = match.group(5)
        numero_cond = match.group(6)
        tem_oficio = match.group(7)
        data_oficio = match.group(8)
        data_cumprimento = match.group(9)
        condicionante = Condicionante(descricao, prazo, cumprida, tempo_restante, situacao, numero_cond, tem_oficio, data_oficio, data_cumprimento)
        condicionantes.append(condicionante)

    return Licenca(numero, tipo, requerente, cnpj, numero_processo, endereco, atividade, classe, porte, potencial_poluidor, coordenadas, validade, condicionantes, data_carimbo)


def get_validade_em_dias(validade):
    # "validade" é uma string nesse formato: 1460 (MIL QUATROCENTOS E SESSENTA) DIAS
    if validade is None:
        return None
    # Extrai o número de dias da string
    dias = int(re.search(r'\d+', validade).group())
    return dias

def calcula_tempo_restante(licenca):
    if licenca is None or licenca.validade is None:
        return None
    validade_em_dias = get_validade_em_dias(licenca.validade)
    if validade_em_dias is None:
        return None
    try:
        data_carimbo = datetime.strptime(licenca.data_carimbo + " 00:00:00", '%d/%m/%Y %H:%M:%S')
    except ValueError:
        # print("Data de carimbo inválida")
        return None
    data_validade = data_carimbo + timedelta(days=validade_em_dias)
    tempo_restante = data_validade - datetime.now()
    dias_restantes = tempo_restante.days
    # print("Dias restantes: ", dias_restantes)
    return dias_restantes



from datetime import datetime, timedelta
import re

def tempo_restante_condicionantes(licenca: Licenca, file_id):
    # verificar se a licença está em outros dicionários que não seja o de arquivos_entregues
    # se tiver em algum outro, retorna a lista de condicionantes sem alterar nada
# unique
    if file_id in Config.ARQUIVOS_ENTREGUES:
        # Calcular o tempo restante para cumprir o prazo da condicionante
        if (licenca.data_carimbo is not None) and (licenca.validade is not None):
            # data_carimbo = datetime.strptime(licenca.data_carimbo, '%d/%m/%Y')
            # data_validade = datetime.strptime(licenca.validade, '%d/%m/%Y')

            for cond in licenca.condicionantes:
                if cond.cumprida == 'False':
                    prazo_texto = cond.prazo

                    # Caso o prazo seja "Não Definido"
                    if prazo_texto == "Não Definido":
                        cond.tempo_restante = None
                    # Caso o prazo seja no formato MES/ANO
                    elif re.match(r'\d{2}/\d{4}', prazo_texto):
                        ano = prazo_texto.split('/')[1]
                        if len(ano) > 4:
                            ano = ano[:4] + ano[5:]  # Remove o quinto caractere

                        mes = int(prazo_texto.split('/')[0])
                        ano = int(ano)  # Converte o ano para inteiro
                        data_prazo = datetime(ano, mes, 1)
                        cond.tempo_restante = (data_prazo - datetime.now()).days
                    # Caso o prazo seja no formato N dias
                    elif 'dias' in prazo_texto and 'antes' not in prazo_texto:
                        dias = int(re.search(r'\d+', prazo_texto).group())
                        # data_prazo = data_carimbo + timedelta(days=dias)
                        # cond.tempo_restante = prazo - (datetime.now() - datetime.strptime(licenca.data_carimbo, '%d/%m/%Y')).days
                        if (cond.tem_oficio == 'True') and (cond.data_oficio is not None):
                            # Converter cond.data_oficio para o formato desejado "dd/mm/yyyy", se necessário
                            if '-' in cond.data_oficio:  # Verifica se o formato está como "yyyy-mm-dd"
                                cond.data_oficio = datetime.strptime(cond.data_oficio, '%Y-%m-%d').strftime('%d/%m/%Y')
                            
                            # Converter cond.data_oficio para um objeto datetime para calcular o tempo restante
                            data_oficio = datetime.strptime(cond.data_oficio, '%d/%m/%Y')
                            
                            # Calcula o tempo restante
                            cond.tempo_restante = dias - (datetime.now() - data_oficio).days
                        else:
                            # Extrair apenas a parte da data antes de converter
                            data_somente = licenca.data_carimbo.split(' ')[0]
                            cond.tempo_restante = dias - (datetime.now() - datetime.strptime(data_somente, '%d/%m/%Y')).days

                    # Caso o prazo seja N dias antes da data de validade
                    elif 'dias' in prazo_texto and 'antes' in prazo_texto:
                        dias = int(re.search(r'\d+', prazo_texto).group())
                        # data_prazo = data_validade - timedelta(days=dias)
                        # pegar somente o numero inicial de licenca.validade
                        validade_em_dias = int(re.search(r'\d+', licenca.validade).group())
                        # cond.tempo_restante = validade_em_dias - 120 - (datetime.now() - datetime.strptime(licenca.data_carimbo, '%d/%m/%Y')).days
                        # Extrair apenas a parte da data antes de converter
                        data_somente = licenca.data_carimbo.split(' ')[0]
                        cond.tempo_restante = validade_em_dias - 120 - (datetime.now() - datetime.strptime(data_somente, '%d/%m/%Y')).days


                    # print(f"Condicionante: {cond.descricao}")
                    # print(f"Prazo: {cond.prazo}")
                    # print(f"Tempo restante: {cond.tempo_restante}")

    return licenca.condicionantes


# Função para verificar se tem alguma condicionante que ainda não teve o prazo preenchido
def condicionantes_preenchidas(file_id):
    # Recuperar a licença
    licenca_string = Config.ARQUIVOS_NAO_ENTREGUES.get(file_id, {}).get('data') or Config.ARQUIVOS_ENTREGUES.get(file_id, {}).get('data')
    licenca = parse_licenca(licenca_string)

    # print("Licenca = ", licenca.requerente)

    if licenca is None:
        return "Licença não encontrada", 404

    # Verificar se todas as condicionantes possuem o prazo preenchido
    for cond in licenca.condicionantes:
        if cond.prazo is None or cond.prazo == '' or cond.prazo == 'N/A' or cond.prazo == 'None':
            # Se alguma condicionante não tiver o prazo preenchido, retorna False
            # print("Condicao nao preenchida")
            return 'False'

    # Se todas as condicionantes tiverem o prazo preenchido, retorna True
    # print("Condicao preenchida")
    return 'True'


def atualizar_situacao_condicionantes(licenca):
    for cond in licenca.condicionantes:
        if cond.tempo_restante is not None and cond.tempo_restante not in ['N/A', 'None', ' None', '']:
            tempo_restante = int(cond.tempo_restante)
            if cond.prazo is not None and cond.prazo not in ['N/A', 'None', '']:
                if cond.cumprida == 'False' or cond.cumprida is False:
                    if 'dias' in cond.prazo and 'antes' not in cond.prazo:
                        prazo = int(re.search(r'\d+', cond.prazo).group())
                    else:
                        prazo = 121  # valor padrão

                    if prazo > 0 and prazo <= 120:
                        if tempo_restante <= 15 and tempo_restante > 0:
                            cond.situacao = 'vermelho'
                        elif tempo_restante <= 30 and tempo_restante > 15:
                            cond.situacao = 'amarelo'
                        elif tempo_restante > 30:
                            cond.situacao = 'verde'
                        elif tempo_restante <= 0:
                            cond.situacao = 'roxo'
                    elif prazo > 120:
                        if tempo_restante <= 30 and tempo_restante > 0:
                            cond.situacao = 'vermelho'
                        elif tempo_restante <= 60 and tempo_restante > 30:
                            cond.situacao = 'amarelo'
                        elif tempo_restante > 60:
                            cond.situacao = 'verde'
                        elif tempo_restante <= 0:
                            cond.situacao = 'roxo'
                else:
                    cond.situacao = 'cumprida'
        else:
            cond.situacao = 'sem_prazo'
            # cond.cumprida = 'True'

    return licenca.condicionantes

def condicionantes_vencendo(file_id):
    # Recuperar a licença
    licenca_string = Config.ARQUIVOS_NAO_ENTREGUES.get(file_id, {}).get('data') or Config.ARQUIVOS_ENTREGUES.get(file_id, {}).get('data')
    licenca = parse_licenca(licenca_string)
    if licenca is None:
        return "Licença não encontrada", 404
    
    licenca.condicionantes = tempo_restante_condicionantes(licenca, file_id)
    
    # Atualizar a situação das condicionantes
    licenca.condicionantes = atualizar_situacao_condicionantes(licenca)
    salvar_dicionarios()

    situacao = ''
    contador_condicionantes_cumpridas = sum(1 for cond in licenca.condicionantes if cond.situacao == 'cumprida')

    if contador_condicionantes_cumpridas == len(licenca.condicionantes):
        return 'concluida'
    
    # Verificar a situação mais crítica
    for cond in licenca.condicionantes:
        if cond.situacao in ['vermelho', 'amarelo', 'verde', 'roxo', 'sem_prazo']:
            if cond.situacao == 'roxo':
                situacao = 'roxo'
            elif cond.situacao == 'vermelho' and situacao != 'roxo':
                situacao = 'vermelho'
            elif cond.situacao == 'amarelo' and situacao not in ['roxo', 'vermelho']:
                situacao = 'amarelo'
            elif cond.situacao == 'verde' and situacao not in ['roxo', 'vermelho', 'amarelo']:
                situacao = 'verde'
            elif cond.situacao == 'sem_prazo' and situacao not in ['roxo', 'vermelho', 'amarelo', 'verde']:
                situacao = 'concluida'
    
    return situacao




def gerar_relatorio_excel():
    licencas = []
    
    for file_id, file_data in Config.ARQUIVOS_ENTREGUES.items():
        licenca = parse_licenca(file_data.get('data'))
        if licenca:
            licencas.append({
                "Nome da Licença": file_data['nome'].split('.docx')[0],
                "Data de Entrega": file_data['data_carimbo'],
                "Dias Restantes": file_data['tempo_restante']
            })
    
    # Ordenar as licenças por Dias Restantes em ordem crescente
    licencas = sorted(licencas, key=lambda x: x['Dias Restantes'])
    
    # Converter para DataFrame
    df = pd.DataFrame(licencas)
    return df

def salvar_relatorio_mongodb(df):
    global collection

    # Converter DataFrame para Excel em um buffer de memória
    excel_buffer = io.BytesIO()
    df.to_excel(excel_buffer, index=False)
    excel_buffer.seek(0)
    
    # Inserir o arquivo no MongoDB
    collection.insert_one({
        'nome': 'relatorio.xlsx',
        'data': excel_buffer.getvalue()
    })

import os
import smtplib
import pandas as pd
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email import encoders


import os
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from email.mime.image import MIMEImage
import pandas as pd
import io
import smtplib

def enviar_email():
    from_email = os.getenv('EMAIL')
    password = os.getenv('PASSWORD')
    to_email = "coisasdolabcidades@gmail.com"
    subject = "LabCidades - VARGEM ALTA - RELATÓRIO LICENÇAS AMBIENTAIS EM MONITORAMENTO"


    # Recuperar o relatório do MongoDB
    relatorio = collection.find_one(sort=[('_id', -1)])
    excel_data = relatorio['data']
    
    # Gerar o dashboard e salvar no MongoDB
    criar_dashboard()

    # Contar as licenças em cada categoria
    df = pd.read_excel(io.BytesIO(excel_data))
    total_licencas = df.shape[0]
    proximas_vencer = df[df['Dias Restantes'] < 120].shape[0]
    situacao_critica = df[(df['Dias Restantes'] >= 120) & (df['Dias Restantes'] <= 150)].shape[0]
    situacao_controlada = df[df['Dias Restantes'] > 150].shape[0]

    # Corpo do e-mail em HTML
    body = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: 'Arial', sans-serif;
                background-color: #f0f0f0;
                margin: 0;
                padding: 0;
            }}
            .container {{
                max-width: 800px;
                margin: 40px auto;
                padding: 20px;
                background-color: #ffffff;
                border-radius: 10px;
                box-shadow: 0 0 20px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                text-align: center;
                background-color: #00027c;
                color: white;
                padding: 20px;
                border-radius: 10px 10px 0 0;
                font-size: 24px;
            }}
            .content {{
                padding: 20px;
                color: #333333;
                line-height: 1.6;
                text-align: center;
            }}
            .content p {{
                margin: 10px 0;
            }}
            .footer {{
                text-align: center;
                padding: 20px;
                color: #777777;
                background-color: #f7f7f7;
                border-radius: 0 0 10px 10px;
            }}
            .info-box {{
                padding: 15px;
                margin: 15px 0;
                border-radius: 8px;
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            }}
            .info-box.vermelho {{
                background-color: #FFCCCC;
                border-left: 6px solid #FF0000;
            }}
            .info-box.amarelo {{
                background-color: #FFFFCC;
                border-left: 6px solid #FFD700;
            }}
            .info-box.verde {{
                background-color: #CCFFCC;
                border-left: 6px solid #32CD32;
            }}
            .highlight {{
                font-weight: bold;
                color: #d9534f;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h2>Monitoramento de Licenças</h2>
            </div>
            <div class="content">
                <p>Prezado(a) Gestor(a),</p>
                <p>Existem <span class="highlight">{total_licencas}</span> licenças ambientais em monitoramento no município de Vargem Alta.</p>
                
                <div class="info-box vermelho">
                    <p>Existem <span class="highlight">{proximas_vencer}</span> licenças ambientais próximas de vencer e passaram do prazo de renovação (menos de 120 dias).</p>
                </div>
                
                <div class="info-box amarelo">
                    <p>Existem <span class="highlight">{situacao_critica}</span> licenças ambientais com situação crítica de prazo e precisam de atenção para possível renovação (entre 120 e 150 dias).</p>
                </div>
                
                <div class="info-box verde">
                    <p>Existem <span class="highlight">{situacao_controlada}</span> licenças ambientais com situação controlada de prazo (mais de 150 dias).</p>
                </div>
                
                <p>Abaixo, você encontrará o dashboard com gráficos atualizados sobre a situação das licenças ambientais em Vargem Alta:</p>
                <img src="cid:dashboard_image" alt="Dashboard" style="width:100%; max-width:800px; border-radius: 8px; margin-top: 10px;"/>
                <p>Em anexo a este e-mail, você encontrará um relatório excel detalhado com as informações de todas as licenças ambientais em monitoramento.</p>
            </div>
            <div class="footer">
                <p>Atenciosamente,</p>
                <p>LabCidades - UFES</p>
                <img src="cid:image1" alt="LabCidades Logo" style="width:150px; height:auto; margin-top: 10px;"/>
            </div>
        </div>
    </body>
    </html>
    """

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'html'))

    part = MIMEBase('application', 'octet-stream')
    part.set_payload(excel_data)
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment; filename=relatorio.xlsx')
    msg.attach(part)

    dashboard_data = Config.DB.imagens.find_one(sort=[('_id', -1)])['imagem']
    part = MIMEImage(dashboard_data, name="dashboard_relatorio_licencas_ambientais.png")
    part.add_header('Content-ID', '<dashboard_image>')
    msg.attach(part)

    with open("assets/LabCidades.jpg", "rb") as img:
        mime_image = MIMEImage(img.read())
        mime_image.add_header('Content-ID', '<image1>')
        mime_image.add_header('Content-Disposition', 'inline', filename='LabCidades.jpeg')
        msg.attach(mime_image)

    # Enviar o e-mail
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(from_email, password)
    text = msg.as_string()
    server.sendmail(from_email, to_email, text)
    server.quit()

    print("E-mail enviado com sucesso!")


def salvar_dashboard_no_mongodb(imagem):
    Config.DB.imagens.insert_one({'imagem': Binary(imagem)})


matplotlib.use('Agg')  # Usar backend 'Agg' para evitar problemas de GUI

def criar_dashboard():
    # Dados para gráficos
    categorias = ['Não Entregues', 'Em Monitoramento', 'Arquivadas']
    conteudo = [len(Config.ARQUIVOS_NAO_ENTREGUES), len(Config.ARQUIVOS_ENTREGUES), len(Config.ARQUIVOS_ENCERRADOS)]

    categorias_monitoramento = ['< 120 dias', '120-150 dias', '> 150 dias']
    conteudo_monitoramento = [0, 0, 0]

    for file_data in Config.ARQUIVOS_ENTREGUES.values():
        tempo_restante = file_data.get('tempo_restante')
        if tempo_restante:
            try:
                if isinstance(tempo_restante, str):
                    dias_restantes = int(tempo_restante.split()[0])
                else:
                    dias_restantes = int(tempo_restante)

                if dias_restantes < 120:
                    conteudo_monitoramento[0] += 1
                elif 120 <= dias_restantes <= 150:
                    conteudo_monitoramento[1] += 1
                else:
                    conteudo_monitoramento[2] += 1
            except ValueError:
                continue

    def func_format(pct, allvalues):
        absolute = int(pct / 100. * sum(allvalues))
        return f'{absolute} ({pct:.1f}%)'

    fig, axs = plt.subplots(1, 2, figsize=(20, 8))

    axs[0].pie(conteudo, labels=categorias, autopct=lambda pct: func_format(pct, conteudo), colors=['#ff9999','#66b3ff','#99ff99'])
    axs[0].set_title('Panorama Geral das Licenças')

    axs[1].pie(conteudo_monitoramento, labels=categorias_monitoramento, autopct=lambda pct: func_format(pct, conteudo_monitoramento), colors=['#ff9999','#ffcc99','#66b3ff'])
    axs[1].set_title('Situação dos Dias Restantes (Em Monitoramento)')

    plt.tight_layout()

    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    plt.close()

    buffer.seek(0)
    salvar_dashboard_no_mongodb(buffer.getvalue())

    return 'Imagem do dashboard salva no MongoDB com sucesso!'
